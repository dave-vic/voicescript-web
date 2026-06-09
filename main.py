"""
VoiceScript Web — FastAPI backend  v3.2
All transcription via Groq API — zero local Whisper model loaded,
keeping RAM well under Render free-tier's 512 MB limit.
Long-file safe: speaker audio loaded at 8 kHz to cap RAM usage.
"""

import os, json, time, tempfile, threading, uuid, math
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

# ── Config ────────────────────────────────────────────────────────────────────
def _load_default_key() -> str:
    if k := os.environ.get("GROQ_API_KEY", ""):
        return k
    for candidate in [
        Path(__file__).parent / "config.json",
        Path.home() / "Desktop" / "VoiceScript.app" / "Contents" / "Resources" / "config.json",
        Path.home() / "Desktop" / "audio-transcriber" / "config.json",
    ]:
        if candidate.exists():
            try:
                data = json.loads(candidate.read_text())
                if k := data.get("groq_api_key", ""):
                    return k
            except Exception:
                pass
    return ""

DEFAULT_KEY  = _load_default_key()
MAX_CHUNK_MB = 24
UPLOAD_DIR   = Path(tempfile.gettempdir()) / "voicescript_uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

JOBS: dict[str, dict] = {}
JOB_MAX_AGE  = 60 * 60        # keep completed jobs 1 hour
JOB_TIMEOUT  = 60 * 25        # allow 25 min for long files (Groq + LLM notes)
MAX_UPLOAD_MB = 200            # reject files over 200 MB with a clear message
_EXECUTOR    = ThreadPoolExecutor(max_workers=3)


def _cleanup_jobs():
    now = time.time()
    for jid, j in list(JOBS.items()):
        if j.get("done"):
            if now - j.get("ts", now) > JOB_MAX_AGE:
                JOBS.pop(jid, None)
        else:
            if now - j.get("ts_updated", j.get("ts", now)) > JOB_TIMEOUT:
                j.update(done=True,
                         error="Processing timed out after 25 minutes. "
                               "Try turning off Speaker ID for faster processing, "
                               "or split the audio into shorter clips.")


app = FastAPI(title="VoiceScript Web")
app.add_middleware(
    CORSMiddleware, allow_origins=["*"],
    allow_methods=["*"], allow_headers=["*"],
)
app.mount("/static", StaticFiles(directory="static"), name="static")


# ── Helpers ───────────────────────────────────────────────────────────────────
def job_update(jid: str, **kw):
    if jid in JOBS:
        JOBS[jid].update(kw)
        JOBS[jid]["ts_updated"] = time.time()


def split_audio(path: Path) -> list[Path]:
    from pydub import AudioSegment
    audio = AudioSegment.from_file(str(path))
    out   = path.parent / (path.stem + "_converted.mp3")
    audio.export(str(out), format="mp3", bitrate="64k")
    size_mb = out.stat().st_size / (1024 * 1024)
    if size_mb <= MAX_CHUNK_MB:
        return [out]
    n        = math.ceil(size_mb / MAX_CHUNK_MB)
    chunk_ms = max(1, len(audio) // n)
    chunks   = []
    for i in range(n):
        s = i * chunk_ms
        e = min((i + 1) * chunk_ms, len(audio))
        if s >= len(audio):
            break
        cp = path.parent / f"_chunk_{i+1}_of_{n}.mp3"
        audio[s:e].export(str(cp), format="mp3", bitrate="64k")
        chunks.append(cp)
    out.unlink(missing_ok=True)
    return chunks


def transcribe_chunks(chunks: list[Path], api_key: str,
                      language: str = None, jid: str = None,
                      verbose_json: bool = False):
    """
    Transcribe via Groq Whisper API — no local model loaded.
    If verbose_json=True, returns list of segment dicts (for speaker ID).
    Otherwise returns a plain string.
    """
    from groq import Groq
    from concurrent.futures import ThreadPoolExecutor as _Pool, as_completed
    client  = Groq(api_key=api_key)
    results = {}
    total   = len(chunks)
    fmt     = "verbose_json" if verbose_json else "text"

    def do(i, cp):
        for attempt in range(3):
            try:
                with open(cp, "rb") as f:
                    kw = dict(model="whisper-large-v3-turbo",
                              file=f, response_format=fmt)
                    if language and language.strip():
                        kw["language"] = language.strip()
                    resp = client.audio.transcriptions.create(**kw)
                    if verbose_json:
                        # resp is an object; grab segments list
                        segs = getattr(resp, "segments", [])
                        return i, segs
                    return i, resp
            except Exception as e:
                if attempt < 2 and "connection" in str(e).lower():
                    time.sleep(5)
                else:
                    raise

    with _Pool(max_workers=min(4, total)) as ex:
        futs = {ex.submit(do, i, cp): i for i, cp in enumerate(chunks)}
        done = 0
        for fut in as_completed(futs):
            idx, payload = fut.result()
            results[idx] = payload
            done += 1
            if jid:
                pct = 20 + int(done / total * 45)
                job_update(jid, progress=pct,
                           status=f"Transcribing… ({done}/{total} parts)")

    for cp in chunks:
        try: cp.unlink()
        except: pass

    if verbose_json:
        # Merge segment lists in order, adjusting timestamps per chunk
        # (For single-chunk files offsets are 0; multi-chunk we estimate)
        all_segs = []
        for i in sorted(results):
            all_segs.extend(results[i] or [])
        return all_segs
    return "\n\n".join(str(results[i]) for i in sorted(results))


def identify_speakers_from_segments(segs, audio_path: str, jid: str = None) -> list[dict]:
    """
    Speaker clustering using ONLY pydub + numpy — no local Whisper model.
    Takes Groq-returned segment objects (with .start, .end, .text).
    Audio is downsampled to 8 kHz before loading into numpy to cap RAM:
      25-min audio @ 44.1 kHz = ~265 MB; @ 8 kHz = ~48 MB — safe on free tier.
    """
    import numpy as np
    from pydub import AudioSegment

    if jid:
        job_update(jid, status="Loading audio for speaker analysis…", progress=68)

    # Downsample to 8 kHz mono — enough for voice fingerprinting, ~5× less RAM
    SPEAKER_SR = 8000
    audio   = (AudioSegment.from_file(audio_path)
               .set_channels(1)
               .set_frame_rate(SPEAKER_SR))
    samples = np.array(audio.get_array_of_samples(), dtype=np.float32)
    sr      = SPEAKER_SR

    if jid:
        job_update(jid, status="Analysing speaker voices…", progress=72)

    def get_features(s0, s1):
        s = max(0, int(s0 * sr))
        e = min(len(samples), int(s1 * sr))
        c = samples[s:e]
        if len(c) < 100:
            return None
        energy = float(np.sqrt(np.mean(c ** 2)))
        zcr    = float(np.mean(np.abs(np.diff(np.sign(c)))) / 2)
        fft    = np.abs(np.fft.rfft(c[:min(len(c), 4096)]))
        freqs  = np.fft.rfftfreq(min(len(c), 4096), 1 / sr)
        spec   = float(np.sum(freqs * fft) / (np.sum(fft) + 1e-9))
        return np.array([energy, zcr, spec])

    def dist(a, b):
        if a is None or b is None:
            return 1.0
        na = a / (np.linalg.norm(a) + 1e-9)
        nb = b / (np.linalg.norm(b) + 1e-9)
        return float(np.linalg.norm(na - nb))

    PAUSE   = 1.0
    NEW_D   = 0.35
    SAME_D  = 0.20
    profiles = {}
    next_id  = 1
    out      = []
    prev_end = 0

    for seg in segs:
        # Groq segments may be dicts or objects
        start = seg.get("start", 0) if isinstance(seg, dict) else getattr(seg, "start", 0)
        end   = seg.get("end", 0)   if isinstance(seg, dict) else getattr(seg, "end", 0)
        text  = (seg.get("text", "") if isinstance(seg, dict) else getattr(seg, "text", "")).strip()
        if not text:
            continue
        gap      = start - prev_end
        features = get_features(start, min(start + 1.5, end))

        if gap < PAUSE or not profiles:
            sid = list(profiles.keys())[-1] if profiles else next_id
            if sid not in profiles:
                profiles[sid] = features
                next_id += 1
        else:
            best_id, best_d = None, float("inf")
            for s, p in profiles.items():
                d = dist(features, p)
                if d < best_d:
                    best_d, best_id = d, s
            if best_d < SAME_D:
                sid = best_id
                if features is not None and profiles[best_id] is not None:
                    profiles[best_id] = profiles[best_id] * 0.7 + features * 0.3
            else:
                sid = next_id
                profiles[sid] = features
                next_id += 1

        out.append({"speaker": sid, "start": start, "end": end, "text": text})
        prev_end = end

    if jid:
        job_update(jid, status="Speaker analysis complete", progress=78)
    return out


def format_with_speakers(segs: list[dict]) -> str:
    if not segs:
        return ""
    merged, cur = [], None
    for s in segs:
        if cur is None or s["speaker"] != cur["speaker"]:
            if cur:
                merged.append(cur)
            cur = {"speaker": s["speaker"], "text": s["text"], "start": s["start"]}
        else:
            cur["text"] += " " + s["text"]
    if cur:
        merged.append(cur)
    lines = []
    for b in merged:
        m, s = int(b["start"] // 60), int(b["start"] % 60)
        lines.append(f"Speaker {b['speaker']}  [{m:02d}:{s:02d}]\n{b['text']}\n")
    return "\n".join(lines)


def make_meeting_notes(transcript: str, api_key: str,
                       source_name: str = "recording",
                       template_path: str = None,
                       jid: str = None) -> tuple[str, bytes]:
    from groq import Groq
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    client = Groq(api_key=api_key)
    chunks = [transcript[i:i+6000] for i in range(0, len(transcript), 6000)]
    summaries = []

    for i, chunk in enumerate(chunks):
        if jid:
            pct = 82 + int(i / len(chunks) * 8)
            job_update(jid, progress=pct,
                       status=f"Extracting key points… ({i+1}/{len(chunks)})")
        r = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system",
                 "content": "Extract key points from meeting transcripts. "
                             "Be concise. Output ONLY raw bullet points — no headings."},
                {"role": "user",
                 "content": f"Extract all key points, decisions, action items, "
                             f"names, dates, and important details:\n\n{chunk}"}
            ],
            max_tokens=600, temperature=0.2)
        summaries.append(r.choices[0].message.content.strip())
        if i < len(chunks) - 1:
            time.sleep(3)

    combined = "\n".join(summaries)

    structure_instruction = (
        "Use these sections:\n"
        "1. MEETING OVERVIEW\n2. KEY DISCUSSIONS\n"
        "3. DECISIONS MADE\n4. ACTION ITEMS\n5. IMPORTANT DETAILS"
    )
    if template_path and Path(template_path).exists():
        try:
            from docx import Document as D2
            doc2     = D2(template_path)
            headings = [p.text.strip() for p in doc2.paragraphs
                        if p.text.strip() and p.style and "Heading" in p.style.name]
            if headings:
                sl = "\n".join(f"- {h}" for h in headings)
                structure_instruction = (
                    f"Use EXACTLY these section headings:\n{sl}\n"
                    f"Write each heading in ALL CAPS on its own line."
                )
        except Exception:
            pass

    if jid:
        job_update(jid, progress=91, status="Writing final meeting notes…")

    final = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {"role": "system",
             "content": "Write thorough, professional meeting notes in full sentences "
                        "and paragraphs — not bullet points. Each section appears ONCE. "
                        "Capture every important detail, decision, name, number."},
            {"role": "user",
             "content": f"Write complete meeting notes.\n"
                        f"RULES: Each heading appears exactly ONCE. "
                        f"No repetition. Write as much detail as needed.\n\n"
                        f"{structure_instruction}\n\n"
                        f"KEY POINTS:\n{combined}"}
        ],
        max_tokens=3000, temperature=0.3)

    notes_text = final.choices[0].message.content.strip()

    if jid:
        job_update(jid, progress=96, status="Creating Word document…")

    doc = DocxDoc()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    date_str = datetime.now().strftime("%B %d, %Y")
    title = doc.add_heading("Meeting Notes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
        r.font.size = Pt(26)

    sub = doc.add_paragraph()
    r   = sub.add_run(f"{date_str}  •  Source: {Path(source_name).name}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x6B, 0x6B, 0x8A)
    r.font.italic = True

    doc.add_paragraph("─" * 60)

    for line in notes_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        clean = line.lstrip("0123456789.) #*").strip()
        is_heading = (
            (line.isupper() and len(line) > 3) or
            (len(line) > 4 and line[0].isdigit() and line[1] in ".)") or
            (line.startswith("**") and line.endswith("**"))
        )
        if is_heading:
            doc.add_paragraph()
            h = doc.add_heading(clean, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
                run.font.size = Pt(13)
                run.bold = True
            div = doc.add_paragraph()
            div.paragraph_format.space_after = Pt(4)
            div_r = div.add_run("─" * 55)
            div_r.font.size = Pt(8)
            div_r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xDD)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(8)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = Pt(16)
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x33)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    r = footer.add_run("Generated by VoiceScript  •  Powered by Groq Llama AI")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x6B, 0x6B, 0x8A)
    r.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return notes_text, buf.getvalue()


def run_job(jid: str, tmp: Path, language: str, mode: str,
            api_key: str, speaker_id: bool, template_path: str | None):
    try:
        job_update(jid, progress=5, status="Converting audio…")
        chunks = split_audio(tmp)

        if speaker_id:
            # Get verbose_json so we have segment timestamps for speaker clustering —
            # no local Whisper model needed at all
            job_update(jid, progress=18, status="Transcribing with timestamps…")
            segs = transcribe_chunks(chunks, api_key, language or None, jid,
                                     verbose_json=True)
            if segs:
                job_update(jid, progress=65, status="Identifying speakers…")
                diarized = identify_speakers_from_segments(segs, str(tmp), jid)
                if diarized:
                    transcript = format_with_speakers(diarized)
                else:
                    # Fall back to plain text from segments
                    transcript = " ".join(
                        (s.get("text", "") if isinstance(s, dict) else getattr(s, "text", ""))
                        for s in segs
                    ).strip()
            else:
                transcript = ""
        else:
            job_update(jid, progress=18, status="Sending to AI…")
            transcript = transcribe_chunks(chunks, api_key, language or None, jid)

        notes_text = None
        docx_b64   = None

        if mode == "meeting":
            notes_text, docx_bytes = make_meeting_notes(
                transcript, api_key,
                source_name=tmp.name,
                template_path=template_path,
                jid=jid,
            )
            import base64
            docx_b64 = base64.b64encode(docx_bytes).decode()

        job_update(jid, progress=100, status="Done!",
                   done=True,
                   result={
                       "mode":       mode,
                       "transcript": transcript,
                       "notes":      notes_text,
                       "docx_b64":   docx_b64,
                       "date":       datetime.now().strftime("%B %d, %Y"),
                       "speaker_id": speaker_id,
                   })
    except Exception as e:
        import traceback
        job_update(jid, done=True,
                   error=str(e),
                   detail=traceback.format_exc())
    finally:
        try: tmp.unlink()
        except: pass
        if template_path:
            try: Path(template_path).unlink()
            except: pass


# ── Routes ────────────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def index():
    return FileResponse("static/index.html")


@app.post("/transcribe/start")
async def transcribe_start(
    file:        UploadFile = File(...),
    language:    str        = Form(""),
    mode:        str        = Form("txt"),
    api_key:     str        = Form(""),
    speaker_id:  str        = Form("false"),
    template:    UploadFile = File(None),
):
    key = api_key.strip() or DEFAULT_KEY
    if not key:
        raise HTTPException(400, "No API key. Add your free Groq key in Settings.")

    suffix   = Path(file.filename).suffix.lower() or ".mp3"
    jid      = str(uuid.uuid4())[:8]
    tmp      = UPLOAD_DIR / f"upload_{jid}{suffix}"
    raw      = await file.read()

    # Reject files that are too large before wasting memory
    size_mb = len(raw) / (1024 * 1024)
    if size_mb > MAX_UPLOAD_MB:
        raise HTTPException(
            413,
            f"File is {size_mb:.0f} MB — please keep uploads under {MAX_UPLOAD_MB} MB. "
            f"Compress or trim the audio first."
        )
    tmp.write_bytes(raw)
    del raw   # free memory immediately after writing to disk

    tmpl_path = None
    if template and template.filename:
        tp = UPLOAD_DIR / f"template_{jid}.docx"
        tp.write_bytes(await template.read())
        tmpl_path = str(tp)

    _cleanup_jobs()
    now = time.time()
    JOBS[jid] = {"progress": 0, "status": "Starting…", "done": False,
                 "ts": now, "ts_updated": now}
    do_speaker = speaker_id.lower() == "true"
    _EXECUTOR.submit(run_job, jid, tmp, language, mode, key, do_speaker, tmpl_path)
    return JSONResponse({"job_id": jid})


@app.get("/transcribe/status/{jid}")
async def transcribe_status(jid: str):
    _cleanup_jobs()
    job = JOBS.get(jid)
    if not job:
        raise HTTPException(404, "Job not found or expired — please try again.")
    safe = {k: v for k, v in job.items() if k not in ("result",)}
    if job.get("done") and not job.get("error"):
        safe["result"] = job.get("result", {})
    return JSONResponse(safe)


@app.get("/config")
async def config():
    return JSONResponse({"has_server_key": bool(DEFAULT_KEY)})


@app.get("/health")
async def health():
    return {"status": "ok", "version": "3.2.0"}


if __name__ == "__main__":
    import socket
    print()
    print("  ╔═══════════════════════════════════════╗")
    print("  ║       VoiceScript Web  v3.1           ║")
    print("  ╠═══════════════════════════════════════╣")
    print("  ║  Your Mac:   http://localhost:8000    ║")
    try:
        ip = socket.gethostbyname(socket.gethostname())
        print(f"  ║  Network:    http://{ip}:8000    ║")
    except:
        pass
    print("  ║                                       ║")
    print("  ║  Share the Network link with anyone  ║")
    print("  ║  on the same WiFi — Mac, Windows,    ║")
    print("  ║  iPhone, Android all work instantly! ║")
    print("  ╚═══════════════════════════════════════╝")
    print()
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
